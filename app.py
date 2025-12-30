"""
White Character Finder - Backend
Detects hidden/invisible characters in PDFs and DOCX files
"""

import os
import re
import tempfile
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
import pdfplumber
from docx import Document
from docx.shared import RGBColor
import unicodedata

app = Flask(__name__, static_folder='static')
CORS(app)

# Configuration
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max

# Zero-width and invisible Unicode characters
INVISIBLE_CHARS = {
    '\u200B': 'Zero Width Space',
    '\u200C': 'Zero Width Non-Joiner',
    '\u200D': 'Zero Width Joiner',
    '\u200E': 'Left-to-Right Mark',
    '\u200F': 'Right-to-Left Mark',
    '\u2060': 'Word Joiner',
    '\u2061': 'Function Application',
    '\u2062': 'Invisible Times',
    '\u2063': 'Invisible Separator',
    '\u2064': 'Invisible Plus',
    '\uFEFF': 'Zero Width No-Break Space (BOM)',
    '\u00AD': 'Soft Hyphen',
    '\u034F': 'Combining Grapheme Joiner',
    '\u061C': 'Arabic Letter Mark',
    '\u115F': 'Hangul Choseong Filler',
    '\u1160': 'Hangul Jungseong Filler',
    '\u17B4': 'Khmer Vowel Inherent Aq',
    '\u17B5': 'Khmer Vowel Inherent Aa',
    '\u180E': 'Mongolian Vowel Separator',
    '\u2000': 'En Quad',
    '\u2001': 'Em Quad',
    '\u2002': 'En Space',
    '\u2003': 'Em Space',
    '\u2004': 'Three-Per-Em Space',
    '\u2005': 'Four-Per-Em Space',
    '\u2006': 'Six-Per-Em Space',
    '\u2007': 'Figure Space',
    '\u2008': 'Punctuation Space',
    '\u2009': 'Thin Space',
    '\u200A': 'Hair Space',
    '\u202A': 'Left-to-Right Embedding',
    '\u202B': 'Right-to-Left Embedding',
    '\u202C': 'Pop Directional Formatting',
    '\u202D': 'Left-to-Right Override',
    '\u202E': 'Right-to-Left Override',
    '\u2066': 'Left-to-Right Isolate',
    '\u2067': 'Right-to-Left Isolate',
    '\u2068': 'First Strong Isolate',
    '\u2069': 'Pop Directional Isolate',
    '\u3000': 'Ideographic Space',
    '\u3164': 'Hangul Filler',
    '\uFFA0': 'Halfwidth Hangul Filler',
}

# Tag characters (used for invisible watermarks)
TAG_CHARS_START = 0xE0000
TAG_CHARS_END = 0xE007F


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def is_white_color(r, g, b, threshold=250):
    """Check if a color is close to white"""
    return r >= threshold and g >= threshold and b >= threshold


def is_near_invisible_color(r, g, b, bg_r=255, bg_g=255, bg_b=255, threshold=10):
    """Check if text color is too close to background (assumed white)"""
    return (abs(r - bg_r) <= threshold and 
            abs(g - bg_g) <= threshold and 
            abs(b - bg_b) <= threshold)


def find_invisible_chars(text):
    """Find invisible/zero-width characters in text"""
    findings = []
    
    for i, char in enumerate(text):
        # Check known invisible characters
        if char in INVISIBLE_CHARS:
            context_start = max(0, i - 10)
            context_end = min(len(text), i + 10)
            context = text[context_start:context_end]
            findings.append({
                'type': 'invisible_char',
                'char_code': f'U+{ord(char):04X}',
                'char_name': INVISIBLE_CHARS[char],
                'position': i,
                'context': repr(context)
            })
        
        # Check tag characters
        if TAG_CHARS_START <= ord(char) <= TAG_CHARS_END:
            findings.append({
                'type': 'tag_char',
                'char_code': f'U+{ord(char):04X}',
                'char_name': f'Tag Character ({unicodedata.name(char, "Unknown")})',
                'position': i
            })
        
        # Check for other control characters
        if unicodedata.category(char) in ('Cf', 'Cc', 'Co') and char not in INVISIBLE_CHARS:
            if char not in '\n\r\t':  # Exclude common whitespace
                findings.append({
                    'type': 'control_char',
                    'char_code': f'U+{ord(char):04X}',
                    'char_name': unicodedata.name(char, 'Unknown Control Character'),
                    'position': i
                })
    
    return findings


def analyze_pdf(file_path):
    """Analyze a PDF file for hidden/white text"""
    results = {
        'white_text': [],
        'invisible_chars': [],
        'small_text': [],
        'hidden_layers': [],
        'total_pages': 0,
        'suspicious_score': 0
    }
    
    try:
        with pdfplumber.open(file_path) as pdf:
            results['total_pages'] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                # Get all characters with their properties
                chars = page.chars
                
                for char_obj in chars:
                    text = char_obj.get('text', '')
                    if not text:
                        continue
                    
                    # Get font size
                    font_size = char_obj.get('size', 12)
                    
                    # Get color if available (stroking_color or non_stroking_color)
                    # pdfplumber provides colors as tuples
                    non_stroking = char_obj.get('non_stroking_color')
                    stroking = char_obj.get('stroking_color')
                    
                    # Check for white/near-white text
                    is_white = False
                    color_str = None
                    
                    # Check non-stroking color (fill color for text)
                    if non_stroking:
                        if isinstance(non_stroking, (list, tuple)):
                            if len(non_stroking) >= 3:
                                # RGB color
                                r, g, b = non_stroking[0], non_stroking[1], non_stroking[2]
                                # Colors in pdfplumber are 0-1 range, convert to 0-255
                                r_255 = int(r * 255) if r <= 1 else r
                                g_255 = int(g * 255) if g <= 1 else g
                                b_255 = int(b * 255) if b <= 1 else b
                                if is_white_color(r_255, g_255, b_255):
                                    is_white = True
                                    color_str = f'rgb({r_255},{g_255},{b_255})'
                            elif len(non_stroking) == 1:
                                # Grayscale
                                gray = non_stroking[0]
                                gray_255 = int(gray * 255) if gray <= 1 else gray
                                if gray_255 >= 250:
                                    is_white = True
                                    color_str = f'gray({gray_255})'
                        elif isinstance(non_stroking, (int, float)):
                            # Single value grayscale
                            gray = non_stroking
                            gray_255 = int(gray * 255) if gray <= 1 else gray
                            if gray_255 >= 250:
                                is_white = True
                                color_str = f'gray({gray_255})'
                    
                    if is_white:
                        # Aggregate consecutive white characters
                        results['white_text'].append({
                            'page': page_num,
                            'text': text[:100] + ('...' if len(text) > 100 else ''),
                            'color': color_str,
                            'font_size': round(font_size, 2)
                        })
                        results['suspicious_score'] += 2
                    
                    # Check for very small text (potential hiding)
                    if font_size < 2 and text.strip():
                        results['small_text'].append({
                            'page': page_num,
                            'text': text[:100] + ('...' if len(text) > 100 else ''),
                            'font_size': round(font_size, 2)
                        })
                        results['suspicious_score'] += 1
                
                # Get full page text and check for invisible characters
                page_text = page.extract_text() or ""
                invisible = find_invisible_chars(page_text)
                for finding in invisible:
                    finding['page'] = page_num
                    results['invisible_chars'].append(finding)
                    results['suspicious_score'] += 3
                
    except Exception as e:
        results['error'] = str(e)
    
    # Consolidate white text findings (group by page)
    if results['white_text']:
        consolidated = {}
        for item in results['white_text']:
            page = item['page']
            if page not in consolidated:
                consolidated[page] = {
                    'page': page,
                    'text': '',
                    'color': item['color'],
                    'font_size': item['font_size'],
                    'count': 0
                }
            consolidated[page]['text'] += item['text']
            consolidated[page]['count'] += 1
        
        results['white_text'] = []
        for page, data in consolidated.items():
            if data['count'] > 0:
                text = data['text'][:100] + ('...' if len(data['text']) > 100 else '')
                results['white_text'].append({
                    'page': page,
                    'text': text,
                    'color': data['color'],
                    'font_size': data['font_size'],
                    'char_count': data['count']
                })
                # Add to suspicious score based on amount
                results['suspicious_score'] += min(data['count'], 20)
    
    return results


def analyze_docx(file_path):
    """Analyze a DOCX file for hidden/white text"""
    results = {
        'white_text': [],
        'invisible_chars': [],
        'small_text': [],
        'hidden_text': [],
        'total_paragraphs': 0,
        'suspicious_score': 0
    }
    
    try:
        doc = Document(file_path)
        results['total_paragraphs'] = len(doc.paragraphs)
        
        for para_num, para in enumerate(doc.paragraphs, 1):
            full_text = para.text
            
            # Check for invisible characters in full paragraph text
            invisible = find_invisible_chars(full_text)
            for finding in invisible:
                finding['paragraph'] = para_num
                results['invisible_chars'].append(finding)
                results['suspicious_score'] += 3
            
            # Check each run for formatting-based hiding
            for run in para.runs:
                text = run.text
                if not text.strip():
                    continue
                
                font = run.font
                
                # Check for white/near-white text color
                if font.color and font.color.rgb:
                    rgb = font.color.rgb
                    r, g, b = rgb[0], rgb[1], rgb[2]
                    if is_white_color(r, g, b):
                        results['white_text'].append({
                            'paragraph': para_num,
                            'text': text[:100] + ('...' if len(text) > 100 else ''),
                            'color': f'rgb({r},{g},{b})'
                        })
                        results['suspicious_score'] += 10
                
                # Check for very small font size
                if font.size and font.size.pt < 2:
                    results['small_text'].append({
                        'paragraph': para_num,
                        'text': text[:100] + ('...' if len(text) > 100 else ''),
                        'font_size': font.size.pt
                    })
                    results['suspicious_score'] += 5
                
                # Check for hidden text property
                if font.hidden:
                    results['hidden_text'].append({
                        'paragraph': para_num,
                        'text': text[:100] + ('...' if len(text) > 100 else ''),
                        'property': 'hidden=True'
                    })
                    results['suspicious_score'] += 15
        
        # Also check headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        invisible = find_invisible_chars(para.text)
                        for finding in invisible:
                            finding['location'] = 'header'
                            results['invisible_chars'].append(finding)
                            results['suspicious_score'] += 3
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        invisible = find_invisible_chars(para.text)
                        for finding in invisible:
                            finding['location'] = 'footer'
                            results['invisible_chars'].append(finding)
                            results['suspicious_score'] += 3
                            
    except Exception as e:
        results['error'] = str(e)
    
    return results


@app.route('/')
def index():
    return send_from_directory('static', 'index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type. Only PDF and DOCX files are allowed.'}), 400
    
    # Save to temporary file
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}') as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name
    
    try:
        # Analyze based on file type
        if ext == 'pdf':
            results = analyze_pdf(tmp_path)
            results['file_type'] = 'PDF'
        else:
            results = analyze_docx(tmp_path)
            results['file_type'] = 'DOCX'
        
        results['filename'] = filename
        
        # Generate verdict
        score = results['suspicious_score']
        if score == 0:
            results['verdict'] = 'clean'
            results['verdict_message'] = 'No hidden characters or suspicious content detected.'
        elif score < 10:
            results['verdict'] = 'low'
            results['verdict_message'] = 'Minor suspicious elements found. Review recommended.'
        elif score < 30:
            results['verdict'] = 'medium'
            results['verdict_message'] = 'Moderate suspicious content detected. Careful review required.'
        else:
            results['verdict'] = 'high'
            results['verdict_message'] = 'High likelihood of hidden content designed to trick AI systems!'
        
        return jsonify(results)
        
    finally:
        # Clean up temp file
        os.unlink(tmp_path)


if __name__ == '__main__':
    os.makedirs('static', exist_ok=True)
    app.run(debug=True, port=5000)
